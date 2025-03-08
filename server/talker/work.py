from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn


def create_styled_word_doc_from_string(test_data_string, output_file="test_cases.docx"):
    """
    Creates a Word document from test case data with styled headers, borders, and without @ symbols.
    The blue background will fill the entire cell.

    Args:
        test_data_string (str): String containing the test case data.
        output_file (str, optional): Name of the output Word document file.
        Defaults to "test_cases.docx".
    """
    document = Document()

    # Create a custom style for the table headers
    header_style = document.styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.bold = True
    header_style.font.color.rgb = RGBColor(255, 255, 255)


    # Split the input string into lines and process them
    lines = test_data_string.strip().split("\n")
    if not lines:
        print("No data to process")
        return

    # Process the first line as the header and remaining as data rows
    headers = [cell.strip().replace('@', '') for cell in lines[0].split("+") if cell.strip()]
    data_rows = [
        [cell.strip() for cell in line.split("+") if cell.strip()]
        for line in lines[1:]
    ]

    # Create the table with header row
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # This line gives the table a style with border

    # Add headers to the table
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), 'ADD8E6')
        cell._tc.get_or_add_tcPr().append(cell_shading)
        p = cell.paragraphs[0]
        p.text = header
        p.style = header_style
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add data rows to the table
    for row_data in data_rows:
        row_cells = table.add_row().cells
        # Create cells for the new row before we try to set the cell data
        for _ in range(len(headers)):
            row_cells[_]
        for i, cell_data in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            p.text = cell_data
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'

    # Save the Word document
    document.save(output_file)
    print(f"Word document saved to '{output_file}'")


if __name__ == '__main__':
    test_data_string = """
+@Test Case@ +@URL@ +@Request Body@ +@Headers@ +@Expected Status Code@ +@Expected Response Body (Partial)@
+1+ +/api/items/1+ + + +200+ +{"item":{"id":1,"name":"Item 1","description":"Description 1"}}+
+2+ +/api/items/2+ + + +200+ +{"item":{"id":2,...}}+
+3+ +/api/items/3+ + + +404+ +{"error":"Item with id 3 not found"}+
+4+ +/api/items/abc+ + + +404+ +(Expect appropriate error handling)+
+5+ +/api/items/-1+ + + +404+ +(Expect appropriate error handling)+
"""
    create_styled_word_doc_from_string(test_data_string)