import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Length

# Configure the API key for Google Generative AI
genai.configure(api_key="AIzaSyC7W9OoTb7cZ_top7bFMK10sLGFkjeS5zA")
model = genai.GenerativeModel("gemini-1.5-flash")


def add_custom_paragraph(doc, text, left_indent):
    """
    Add a paragraph with custom formatting based on provided parameters.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True  # Make text bold

    # Set indentation
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(left_indent)

def create_styled_table(doc, table_data, header_style):
    """
    Creates a table in the document with styled headers, borders, and precise width.
    """
    if not table_data:
        return
    
    # Page width is typically 8.5 inches, subtract margins
    page_width = Inches(8.5 - (2 * 0.05))  # Subtracting 5px from each side
    
    header_line = table_data[0]
    headers = [part.strip() for part in header_line.split("@") if part.strip()]

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to be slightly less than page width
    table_properties = table._element.xpath('//w:tblPr')[0]
    table_width = OxmlElement('w:tblW')
    # Convert page width to twips (1 inch = 1440 twips)
    table_width.set(qn('w:w'), str(int(page_width * 1440)))
    table_width.set(qn('w:type'), 'dxa')  # Direct width in twips
    table_properties.append(table_width)

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        
        # Add light blue shading to header
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), 'ADD8E6')
        cell._tc.get_or_add_tcPr().append(cell_shading)
        
        # Add border to header cells
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_type in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tc_pr.append(borders)

        # Format header text
        p = cell.paragraphs[0]
        p.text = header
        p.style = header_style
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.bold = True

    # Add data rows
    for row_data in table_data[1:]:
        if row_data.startswith("+"):
            row_parts = [part.strip() for part in row_data.split("+") if part.strip()]
            row_cells = table.add_row().cells
            
            for i, cell_data in enumerate(row_parts):
                if i < len(row_cells):
                    cell = row_cells[i]
                    
                    # Add border to data cells
                    tc_pr = cell._tc.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    for border_type in ['top', 'bottom', 'left', 'right']:
                        border = OxmlElement(f'w:{border_type}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:color'), '000000')
                        borders.append(border)
                    tc_pr.append(borders)

                    # Format cell text
                    p = cell.paragraphs[0]
                    p.text = cell_data
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.size = Pt(10)

    # Distribute table columns evenly
    num_columns = len(headers)
    column_width = page_width / num_columns
    for column in table.columns:
        column.width = Inches(column_width)

    return table

def process_plain_text_to_word(plain_text, output_file="structured_document.docx"):
    """
    Process plain text into a structured Word document with proper formatting.
    """
    try:
        print(f"Received text for processing:\n{plain_text}")

        # Prepare the prompt for the Generative AI
        prompt = (
            "Analyze the following plain text and organize it into a structured document with appropriate headings like proper headings,sub headings,sub-sub headings,good paragraphs\n"
            "Follow the given seven rules strictly all and generate a document with the structured content.\n"
            "rule1: Actual purpose of the report or the Actual Title should start with % and ends with %\n"
            "rule2: Heading should start with # and ends with # and also maintain indexing like 1,2,3,..\n"
            "rule3: Subheading should start with $ and ends with $ and also maintain indexing like 1.1,1.2,1.3,..\n"
            "rule4: Sub-Subheading should start with * and ends with * and also maintain indexing like a,b,c..\n"
            "rule5: Paragraph should start with -- and ends with --\n"
            "for creating tables use this structure\n"
            "rule6: Table headers should start with @ and ends with @. Parse the text between these symbols as headers and the text before the second @ symbol will be the cell value for the first header next text between the third and fourth @ symbol will be for second header,etc. (these headers should be present in the same row). \n"
            "rule7: Table rows should start with + and ends with + like +value1+,+value2+,.. (this will be considered as table data rows) and parse cell values based on the symbol + \n"
            f"\n{plain_text}\n"
            "Output the structure in a hierarchical format (e.g., numbered headings, subheadings, and paragraphs  and also use symbols %,#,$,*,--,@,+ and remaining symbols must be striclty prohibited for any use  and also this symbols must used for rule defining only).\n"
            "exception if symbols like / are used to specify the path then they can be used for that purpose only like C:/drive/.. path specifying\n"
        )
        print(f"Prompt for the AI model:\n{prompt}")

        # Generate the structured text using the AI model
        response = model.generate_content(prompt)
        if not hasattr(response, 'text'):
            raise ValueError("Invalid response from the AI model")

        structured_text = response.text
        print(f"Structured response:\n{structured_text}")

        # Create a Word document
        doc = Document()

        # Create the header style here
        header_style = doc.styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.bold = True
        header_style.font.color.rgb = RGBColor(255, 255, 255)

        # Apply a light blue background to the headers
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'ADD8E6')  # Light blue color
        header_style.element.get_or_add_pPr().append(shading_elm)

        # Set page margins (optional - you can adjust these values)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # Get the section's properties
            section_properties = section._sectPr

            # Create border elements for all sides
            borders_xml = parse_xml(f'''
                <w:pgBorders {nsdecls('w')}>
                    <w:top w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                </w:pgBorders>
            ''')

            # Add borders
            section_properties.append(borders_xml)

        lines = structured_text.split("\n")
        table_data = []
        in_table = False

        for line in lines:
            if line.startswith("%") and line.endswith("%"):
                content = line.strip('%')
                add_custom_paragraph(doc, content, left_indent=3)
            elif line.startswith("#") and line.endswith("#"):
                content = line.strip('#')
                add_custom_paragraph(doc, content, left_indent=9)
            elif line.startswith("$") and line.endswith("$"):
                content = line.strip('$')
                add_custom_paragraph(doc, content, left_indent=12)
            elif line.startswith("*") and line.endswith("*"):
                content = line.strip('*')
                add_custom_paragraph(doc, content, left_indent=18)
            elif line.startswith("@"):
              table_data.append(line)
              in_table = True
            elif line.startswith("+") and in_table:
                table_data.append(line)
            elif not (line.startswith("@")) and not(line.startswith("+")) and in_table:
                create_styled_table(doc, table_data, header_style)
                table_data = []
                in_table = False
                 # Regular paragraph
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.left_indent = Pt(36)  # Default paragraph indentation
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.left_indent = Pt(36)  # Default paragraph indentation
        if in_table:
           create_styled_table(doc, table_data, header_style)

        # Save the document
        doc.save(output_file)
        print(f"Document saved as {output_file}")
        return output_file

    except Exception as e:
        print(f"Error: {str(e)}")
        return None