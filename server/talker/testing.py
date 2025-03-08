import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Configure the API key for Google Generative AI
genai.configure(api_key="")
model = genai.GenerativeModel("gemini-1.5-flash")

def add_custom_paragraph(doc, text, left_indent):
    """
    Add a paragraph with custom formatting based on provided parameters.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True  # Make text bold

    # Set indentation
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(left_indent)

def process_plain_text_to_word(plain_text, output_file="structured_document.docx"):
    """
    Process plain text into a structured Word document with proper formatting.
    """
    try:
        print(f"Received text for processing:\n{plain_text}")

        # Prepare the prompt for the Generative AI
        prompt = (
            "Analyze the following plain text and organize it into a structured document with appropriate headings like proper headings,sub headings,sub-sub headings,good paragraphs\n"
            "Follow the given five rules strictly all and generate a document with the structured content.\n"
            "rule1: Actual purpose of the report or the Actual Title should start with % and ends with %\n"
            "rule2: Heading should start with # and ends with # and also maintain indexing like 1,2,3,..\n"
            "rule3: Subheading should start with $ and ends with $ and also maintain indexing like 1.1,1.2,1.3,..\n"
            "rule4: Sub-Subheading should start with * and ends with * and also maintain indexing like a,b,c..\n"
            "rule5: Paragraph should start with -- and ends with --\n"
            f"\n{plain_text}\n"
            "Output the structure in a hierarchical format (e.g., numbered headings, subheadings, and paragraphs)."
        )

        # Generate the structured text using the AI model
        response = model.generate_content(prompt)
        if not hasattr(response, 'text'):
            raise ValueError("Invalid response from the AI model")

        structured_text = response.text
        print(f"Structured response:\n{structured_text}")

        # Create a Word document
        doc = Document()
            # Set page margins (optional - you can adjust these values)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
            # Get the section's properties
            section_properties = section._sectPr
        
            # Create border elements for all sides
            borders_xml = parse_xml(f'''
                <w:pgBorders {nsdecls('w')}>
                    <w:top w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="24" w:space="0" w:color="000000"/>
                </w:pgBorders>
            ''')
        
            # Add borders
            section_properties.append(borders_xml)
            
        lines = structured_text.split("\n")
        for line in lines:
            if line.startswith("%") and line.endswith("%"):
                content = line.strip('%')
                add_custom_paragraph(doc, content, left_indent=3)
            elif line.startswith("#") and line.endswith("#"):
                content = line.strip('#')
                add_custom_paragraph(doc, content, left_indent=9)
            elif line.startswith("$") and line.endswith("$"):
                content = line.strip('$')
                add_custom_paragraph(doc, content, left_indent=12)
            elif line.startswith("*") and line.endswith("*"):
                content = line.strip('*')
                add_custom_paragraph(doc, content, left_indent=18)
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.left_indent = Pt(36)  # Default paragraph indentation

        # Save the document
        doc.save(output_file)
        print(f"Document saved as {output_file}")
        return output_file

    except Exception as e:
        print(f"Error: {str(e)}")
        return None

# Example usage
if __name__ == "__main__":
    input_text = """## Report on Python Code for Structuring Plain Text into a Word Document This report analyzes the provided Python code designed to process plain text and generate a structured Word document using Google Generative AI. The analysis covers required imports, module usage, logical and syntactical errors, and an overall code rating. **1. Required Imports and Installation Commands:** The code utilizes several Python libraries. The installation commands are as follows: * **`python -m pip install python-docx`**: Installs the `python-docx` library for Word document manipulation. * **`python -m pip install google-generativeai`**: Installs the Google Generative AI library. The `os` module is a standard Python library and doesn't require separate installation. **2. Python Modules Used and Their Logical Approach:** The code uses the following modules: * **`os`:** While imported, it's not used in the provided code. This suggests unnecessary import. * **`python-docx`:** This library is used for creating and manipulating Word documents. The code leverages its capabilities to add paragraphs, format text (font size, bold), adjust indentation, and add borders using OXML elements for fine-grained control over paragraph properties. The logical approach is to build the document programmatically, adding elements and formatting as needed. * **`google.generativeai`:** This library interacts with Google's Generative AI models. The code uses it to send a prompt to the `gemini-1.5-flash` model, aiming to structure the input plain text into a hierarchical format with headings, subheadings, and paragraphs based on specified rules. The logical approach is to delegate the complex task of text structuring to the AI and then process the AI's output to create the Word document. * **`OxmlElement` and `qn` from `docx.oxml`:** These are used for low-level manipulation of the Word document's XML structure, allowing the addition of borders to paragraphs. This demonstrates a deeper understanding of the `python-docx` library and its capabilities beyond high-level functions. The overall logical approach is to: 1. Receive plain text input. 2. Prepare a prompt for the Google Generative AI model, outlining the desired structure and rules. 3. Send the prompt to the AI model and receive a structured response. 4. Parse the AI's response, identifying headings, subheadings, and paragraphs based on specific markers ("Heading:", "Subheading:", etc.). 5. Create a Word document and add the parsed elements, applying appropriate formatting (font size, bold, indentation, borders) based on their type and level. 6. Save and return the generated Word document. **3. Errors in the Logical Approach:** * **Over-reliance on AI for Structuring:** The code heavily relies on the AI to correctly interpret and structure the input text. The success of the code is entirely dependent on the AI's ability to understand and follow the provided rules. This makes the code brittle and prone to failure if the AI's output is not in the expected format. Robust error handling is needed to gracefully handle unexpected AI responses. * **Rigid Structure Parsing:** The code assumes that the AI will always output text with "Heading:", "Subheading:", etc. prefixes. A more flexible parsing mechanism should be implemented to handle variations in the AI's output format. For example, using regular expressions or other more sophisticated parsing techniques would improve robustness. * **Lack of Input Validation:** The code lacks validation of the input plain text. It assumes the input will always be suitable for AI processing. Adding input validation would improve the code's reliability and prevent unexpected errors. * **Limited Error Handling:** While a `try...except` block is used, it's too general. More specific exception handling is needed to catch and handle different types of errors (e.g., AI API errors, file I/O errors, parsing errors). * **Hardcoded Formatting:** The indentation values (Pt(36), Pt(72), Pt(108)) are hardcoded. It would be better to define these as constants or calculate them dynamically based on a consistent unit (e.g., inches or centimeters). **4. Errors in Syntax and Suggestions for Better Approaches:** The syntax is largely correct, but improvements are possible: * **Unnecessary `os` import:** The `os` module is imported but not used. It should be removed. * **API Key Hardcoding:** The API key is hardcoded directly in the script. This is a security risk. Best practice is to store the API key in environment variables. * **Inconsistent Indentation:** While the code generally follows consistent indentation, some lines within the `add_bordered_paragraph` function could be better formatted for readability. * **Improved Prompt Engineering:** The prompt could be significantly improved. The rules are very rigid, and the AI may struggle to adhere to them perfectly, especially in cases of complex or ambiguous input. A more flexible and natural language prompt, possibly specifying the desired output structure in a more descriptive way, could lead to better results. For example, instead of rigid rules, the prompt could say: "Organize this text into a hierarchical structure with headings, subheadings, and paragraphs. Use a clear, numbered structure for headings and subheadings". * **Enhanced Error Handling:** The `try...except` block should be more specific. It should catch potential `google.generativeai` errors, `python-docx` errors, and `IOError` exceptions separately to provide more informative error messages and potentially handle them differently. **5. Code Rating (out of 10):** The code is functional but has significant limitations in terms of robustness, error handling, and maintainability. We will consider the following metrics: * **Functionality (3/5):** The code achieves its basic functionality, but its dependence on the AI makes it unreliable. * **Readability (4/5):** The code is generally well-structured and uses meaningful variable names, but some areas could be improved for clarity. * **Robustness (2/5):** The code lacks robust error handling and is sensitive to variations in the AI's output. * **Maintainability (3/5):** The code is relatively easy to understand, but improvements in error handling and input validation would significantly enhance its maintainability. * **Security (1/5):** Hardcoding the API key is a major security flaw. **Overall Rating: 2.6/10** The code demonstrates some basic understanding of both the `python-docx` library and Google's Generative AI, but significant improvements are required to address the identified logical and practical flaws before it can be considered a robust and reliable solution. The heavy reliance on an external AI service, coupled with inadequate error handling and security considerations, significantly lowers its overall score."""
    output_file_name = process_plain_text_to_word(input_text)
    if output_file_name:
        print(f"Word document created successfully: {output_file_name}")
    else:
        print("Failed to create Word document.")
